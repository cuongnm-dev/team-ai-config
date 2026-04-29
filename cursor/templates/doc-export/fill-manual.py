#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
fill-manual.py — Fill huong-dan-su-dung.docx template (ETC format)

Strategy:
  1. Replace cover + header/footer placeholder text (rebuild runs, split-run safe)
  2. Fill Section I (Tổng quan tài liệu) — Mục đích, Phạm vi, Định nghĩa, Tài liệu liên quan
  3. Fill revision history table + signing pages (placeholder markers for human completion)
  4. Find "NỘI DUNG" heading in XML body → delete everything after
  5. Rebuild Section II: Giới thiệu chung, Giới thiệu các chức năng, HDSD step-by-step
  6. Rebuild Section 4: Các vấn đề thường gặp (from error-cases)
  7. Mark TOC as dirty so Word refreshes on open

Usage:
  python fill-manual.py \
    --template        /path/docs/templates/huong-dan-su-dung.docx \
    --flow-report     /path/docs/generated/intel/flow-report.json \
    --screenshot-map  /path/docs/generated/intel/screenshot-map.json \
    --screenshots-dir /path/docs/generated/screenshots \
    --output-dir      /path/docs/generated/output \
    --project-name    "Tên Dự Án" \
    --client-name     "Tên Khách Hàng" \
    --dev-unit        "Tên Đơn Vị Phát Triển"
"""

import argparse
import json
import os
import re
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CONTENT_STYLE = "ETC_Content"
LIST_STYLE = "List Paragraph"
PLACEHOLDER_MARK = "[CẦN BỔ SUNG"  # marker for human reviewer


# ── Format normalizers ────────────────────────────────────────────────────────

def normalize_service(service, project_name):
    return {
        "name": service.get("name") or service.get("service-name", "he-thong"),
        "display-name": service.get("display-name") or service.get("name") or project_name,
        "features": service.get("features", []),
    }


def normalize_steps(feature):
    """Support both new steps[] and legacy main-flow[] strings."""
    if feature.get("steps"):
        return feature["steps"]
    steps = []
    for item in feature.get("main-flow", []):
        m = re.match(r"^(\d+)\.\s+(.+)", item.strip())
        if m:
            steps.append({"no": int(m.group(1)), "action": m.group(2), "expected": ""})
        else:
            steps.append({"no": len(steps) + 1, "action": item.strip(), "expected": ""})
    return steps


def normalize_error_cases(feature):
    """Support both new object format {trigger-step, condition, message} and legacy string format."""
    cases = feature.get("error-cases", [])
    if not cases:
        return []
    if isinstance(cases[0], str):
        return [{"trigger-step": None, "condition": c, "message": c} for c in cases]
    return cases


def load_screenshot_map(path):
    """Load screenshot-map.json v2.0 (flat by feature-id). Supports legacy format."""
    if not os.path.exists(path):
        return {}
    with open(path, encoding="utf-8") as f:
        raw = json.load(f)
    # v2.0 schema: {version: "2.0", map: {...}}
    if raw.get("version") == "2.0" and "map" in raw:
        return raw["map"]
    # Legacy flat: {feat_id: [items]}
    if isinstance(raw, dict) and "screenshots" not in raw and "total" not in raw:
        # Wrap legacy flat into v2 structure for consistent downstream access
        return {k: {"screenshots": v if isinstance(v, list) else v.get("screenshots", [])}
                for k, v in raw.items()}
    # Legacy v1: {total, screenshots: [{feature-id, ...}]}
    index = {}
    for sc in raw.get("screenshots", []):
        fid = sc.get("feature-id")
        if fid:
            index.setdefault(fid, {"screenshots": []})["screenshots"].append({
                "file": sc.get("file", ""),
                "step-no": sc.get("step-no"),
                "state": sc.get("state", ""),
                "description": sc.get("description", ""),
            })
    return index


# ── Text replacement (split-run safe) ─────────────────────────────────────────

def _replace_para_text(para, old, new):
    """Replace text in paragraph — rebuild runs to avoid split-run bug."""
    if old not in para.text:
        return False
    new_full = para.text.replace(old, new)
    if not para.runs:
        return False
    para.runs[0].text = new_full
    for run in para.runs[1:]:
        run.text = ""
    return True


def _iter_all_paragraphs(doc):
    """Iterate ALL paragraphs including headers, footers, table cells, nested tables."""
    # Body paragraphs
    for para in doc.paragraphs:
        yield para
    # Body tables (including nested)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
                for nested_table in cell.tables:
                    for nrow in nested_table.rows:
                        for ncell in nrow.cells:
                            for npara in ncell.paragraphs:
                                yield npara
    # Headers and footers (FIX: previously missed)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                for para in header.paragraphs:
                    yield para
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                yield para
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                for para in footer.paragraphs:
                    yield para


def replace_placeholders(doc, project_name, client_name, dev_unit, today):
    """Replace all cover/header/footer placeholders."""
    replacements = [
        ("<Tên dự án>", project_name),
        ("Ban hành: dd/mm/yyyy", f"Ban hành: {today}"),
        ("HỆ THỐNG QUẢN LÝ", project_name.upper()),
        # Header company name: replace with dev-unit (đơn vị phát triển)
        ("CÔNG TY CỔ PHẦN HỆ THỐNG CÔNG NGHỆ ETC", dev_unit.upper()),
    ]
    for para in _iter_all_paragraphs(doc):
        for old, new in replacements:
            _replace_para_text(para, old, new)


# ── XML navigation helpers ────────────────────────────────────────────────────

def _get_para_style(elem):
    ps = elem.find(f".//{{{WNS}}}pStyle")
    return ps.get(f"{{{WNS}}}val", "") if ps is not None else ""


def _get_para_text(elem):
    return "".join(t.text or "" for t in elem.findall(f".//{{{WNS}}}t"))


def _find_heading_by_text(doc, text_needle, style_needle="HEADING"):
    """Find XML paragraph element by heading text (case-insensitive contains)."""
    body = doc.element.body
    for i, child in enumerate(list(body)):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            style = _get_para_style(child).upper()
            text = _get_para_text(child).upper()
            if style_needle.upper() in style and text_needle.upper() in text:
                return i, child
    return None, None


def clear_after_heading(doc, heading_text):
    """Delete all elements after a heading (except trailing sectPr)."""
    body = doc.element.body
    children = list(body)
    idx, _ = _find_heading_by_text(doc, heading_text)
    if idx is None:
        return False
    to_remove = []
    for child in children[idx + 1:]:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag != "sectPr":
            to_remove.append(child)
    for child in to_remove:
        body.remove(child)
    return True


def mark_toc_dirty(doc):
    """Mark TOC fields as dirty so Word refreshes on open (auto-update page numbers)."""
    settings = doc.settings.element
    ns = f"{{{WNS}}}"
    # Set updateFields = true
    update_fields = settings.find(f"{ns}updateFields")
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    # Also mark all TOC fieldChar as dirty
    for fld_char in doc.element.body.iter(f"{ns}fldChar"):
        if fld_char.get(qn("w:fldCharType")) == "begin":
            fld_char.set(qn("w:dirty"), "true")


# ── Paragraph / Table builders ────────────────────────────────────────────────

def add_para(doc, text, style_name, bold=False, color=None):
    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if style_name in (CONTENT_STYLE, LIST_STYLE):
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_placeholder_para(doc, text):
    """Add paragraph marked as needing human review (gray italic)."""
    try:
        p = doc.add_paragraph(style=CONTENT_STYLE)
    except KeyError:
        p = doc.add_paragraph()
    run = p.add_run(f"{PLACEHOLDER_MARK}: {text}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    return p


def _set_table_full_width(table):
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "5000")
    tbl_w.set(qn("w:type"), "pct")


def add_image(doc, screenshots_dir, img_filename, caption=""):
    full_path = os.path.join(screenshots_dir, img_filename)
    if os.path.exists(full_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(full_path, width=Inches(5.5))
        if caption:
            cap_p = add_para(doc, caption, CONTENT_STYLE)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cap_p.runs:
                r.italic = True
                r.font.size = Pt(10)
    else:
        add_placeholder_para(doc, f"Hình: {caption or img_filename} chưa có (screenshot capture thất bại)")


# ── Section I: Tổng quan tài liệu (Mục đích, Phạm vi, Định nghĩa, Tài liệu liên quan) ──

def fill_section_i(doc, project_name, flow, stack):
    """Fill Section I — placeholder replacement within existing template structure.

    Strategy: find each placeholder bracket [Phần này...] and replace with generated content.
    This preserves template formatting (headings, style) — only swaps body text.
    """
    services = flow.get("services", [])
    all_features = [f for svc in services for f in svc.get("features", [])]
    total_features = len(all_features)
    actors = sorted({a for f in all_features for a in f.get("actors", [])})
    stacks_list = stack.get("stacks", []) if stack else []
    framework_names = ", ".join(s.get("framework", "") for s in stacks_list if s.get("framework"))

    # 1. Mục đích tài liệu
    muc_dich_content = (
        f"Tài liệu này được xây dựng nhằm hướng dẫn chi tiết cách sử dụng các chức năng "
        f"của hệ thống {project_name}. Nội dung tài liệu trình bày theo trình tự từng chức năng, "
        f"kèm hình ảnh minh họa và các bước thực hiện cụ thể, giúp người dùng dễ dàng thao tác trên hệ thống. "
        f"Tài liệu này được cung cấp cho Phần mềm phiên bản 1.0."
    )

    # 2. Phạm vi tài liệu
    features_bullets = "\n".join(
        f"- {svc.get('display-name', svc.get('name', ''))}"
        for svc in services
    )
    pham_vi_content = (
        f"Tài liệu này áp dụng cho hệ thống {project_name} với {total_features} chức năng "
        f"được tổ chức theo các phân hệ sau:\n{features_bullets}\n\n"
        f"Đối tượng sử dụng tài liệu: {', '.join(actors) if actors else '[CẦN BỔ SUNG: danh sách người dùng]'}."
    )

    # Replace placeholders in Section I paragraphs
    body = doc.element.body
    in_section_i = False
    i = 0
    children = list(body)
    while i < len(children):
        child = children[i]
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            style = _get_para_style(child).upper()
            text = _get_para_text(child).upper()
            # Detect Section I entry/exit
            if "HEADING" in style and "TỔNG QUAN TÀI LIỆU" in text:
                in_section_i = True
            elif "HEADING" in style and "NỘI DUNG" in text and in_section_i:
                in_section_i = False
                break
            # Within Section I: detect placeholder [Phần này trình bày...] and <Ví dụ:...>
            if in_section_i:
                para_text = _get_para_text(child)
                # Detect placeholder pattern — starts with [ and contains "Phần này"
                if para_text.strip().startswith("[") and "Phần này" in para_text:
                    # Determine which subsection by looking at previous heading
                    subsection = _find_previous_heading(children, i)
                    replacement = None
                    if "Mục đích" in subsection:
                        replacement = muc_dich_content
                    elif "Phạm vi" in subsection:
                        replacement = pham_vi_content
                    if replacement:
                        _replace_para_full_text(child, replacement)
        i += 1

    # 3. Định nghĩa, thuật ngữ — fill table
    _fill_definition_table(doc, project_name, framework_names)

    # 4. Tài liệu liên quan — fill table
    _fill_related_docs_table(doc)


def _find_previous_heading(children, current_idx):
    """Walk back from current to find nearest heading paragraph text."""
    for j in range(current_idx - 1, -1, -1):
        prev = children[j]
        tag = prev.tag.split("}")[-1] if "}" in prev.tag else prev.tag
        if tag == "p":
            style = _get_para_style(prev).upper()
            if "HEADING" in style:
                return _get_para_text(prev)
    return ""


def _replace_para_full_text(para_elem, new_text):
    """Replace entire paragraph text while preserving first run's formatting."""
    runs = para_elem.findall(f".//{{{WNS}}}r")
    if not runs:
        return
    # Keep first run, set its text; clear others
    first_run = runs[0]
    # Remove all w:t children in first run, add new one
    for t in first_run.findall(f"{{{WNS}}}t"):
        first_run.remove(t)
    new_t = OxmlElement("w:t")
    new_t.text = new_text
    new_t.set(qn("xml:space"), "preserve")
    first_run.append(new_t)
    # Clear other runs
    for run in runs[1:]:
        for t in run.findall(f"{{{WNS}}}t"):
            t.text = ""


def _fill_definition_table(doc, project_name, framework_names):
    """Find 'Định nghĩa, thuật ngữ' table and fill with basic terms."""
    terms = [
        ("HDSD", "Hướng dẫn sử dụng", "Tài liệu mô tả cách thao tác phần mềm"),
        ("UI", "User Interface", "Giao diện người dùng"),
        ("API", "Application Programming Interface", "Giao diện lập trình ứng dụng"),
        ("JWT", "JSON Web Token", "Chuẩn mã thông báo truy cập xác thực người dùng"),
    ]
    if framework_names:
        terms.append((framework_names[:15], framework_names, "Công nghệ nền tảng của hệ thống"))

    # Find table under heading "Định nghĩa, thuật ngữ"
    body = doc.element.body
    children = list(body)
    target_idx = None
    for i, child in enumerate(children):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            text = _get_para_text(child)
            if "Định nghĩa" in text and "thuật ngữ" in text:
                target_idx = i
                break
    if target_idx is None:
        return

    # Find first table after target_idx
    for j in range(target_idx + 1, len(children)):
        tag = children[j].tag.split("}")[-1] if "}" in children[j].tag else children[j].tag
        if tag == "tbl":
            table = [t for t in doc.tables if t._tbl == children[j]]
            if table:
                _populate_terms_table(table[0], terms)
            break


def _populate_terms_table(table, terms):
    """Populate (or append) rows in definition table."""
    # Skip header row (row 0)
    existing_data_rows = len(table.rows) - 1
    for idx, (short, full, explanation) in enumerate(terms):
        if idx + 1 < len(table.rows):
            row = table.rows[idx + 1]
        else:
            row = table.add_row()
        cells = row.cells
        if len(cells) >= 3:
            cells[0].text = short
            cells[1].text = full
            cells[2].text = explanation


def _fill_related_docs_table(doc):
    """Placeholder for related docs — leave for human to fill."""
    body = doc.element.body
    children = list(body)
    target_idx = None
    for i, child in enumerate(children):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            text = _get_para_text(child)
            if "Tài liệu liên quan" in text:
                target_idx = i
                break
    if target_idx is None:
        return
    # Find first table — fill row 1 with placeholder
    for j in range(target_idx + 1, min(target_idx + 5, len(children))):
        tag = children[j].tag.split("}")[-1] if "}" in children[j].tag else children[j].tag
        if tag == "tbl":
            table = [t for t in doc.tables if t._tbl == children[j]]
            if table and len(table[0].rows) >= 2:
                row = table[0].rows[1]
                if len(row.cells) >= 3:
                    row.cells[0].text = "1"
                    row.cells[1].text = f"{PLACEHOLDER_MARK}: Tài liệu thiết kế cơ sở]"
                    row.cells[2].text = f"{PLACEHOLDER_MARK}: Đường dẫn/mã tài liệu]"
            break


# ── Revision history + signing pages ──────────────────────────────────────────

def fill_revision_history(doc, today):
    """Find 'CÁC THAY ĐỔI' table and fill first data row."""
    body = doc.element.body
    children = list(body)
    target_idx = None
    for i, child in enumerate(children):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            text = _get_para_text(child)
            if "CÁC THAY ĐỔI" in text:
                target_idx = i
                break
    if target_idx is None:
        return
    for j in range(target_idx + 1, min(target_idx + 5, len(children))):
        tag = children[j].tag.split("}")[-1] if "}" in children[j].tag else children[j].tag
        if tag == "tbl":
            tables = [t for t in doc.tables if t._tbl == children[j]]
            if tables and len(tables[0].rows) >= 2:
                row = tables[0].rows[1]
                cells = row.cells
                if len(cells) >= 5:
                    cells[0].text = today
                    cells[1].text = "Toàn bộ tài liệu"
                    cells[2].text = "A"
                    cells[3].text = "Tạo mới tài liệu"
                    cells[4].text = "1.0"
            break


def fill_signing_pages(doc, dev_unit, client_name):
    """Fill '<Họ tên>' and '<Chức danh>' placeholders on signing pages."""
    replacements = [
        ("<Họ tên>", f"{PLACEHOLDER_MARK}: Tên người ký]"),
        ("<Chức danh>", f"{PLACEHOLDER_MARK}: Chức danh]"),
    ]
    for para in _iter_all_paragraphs(doc):
        for old, new in replacements:
            _replace_para_text(para, old, new)


# ── Section II + Section 4 rebuild ────────────────────────────────────────────

def add_feature_table(doc, features):
    """Create feature catalog table: STT | Chức năng | Mô tả | Đối tượng sử dụng."""
    if not features:
        return
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    _set_table_full_width(table)
    hdr = table.rows[0].cells
    for i, title in enumerate(["STT", "Chức năng", "Mô tả", "Đối tượng sử dụng"]):
        hdr[i].text = title
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    for i, feat in enumerate(features):
        row = table.add_row().cells
        row[0].text = str(i + 1)
        row[1].text = feat.get("name", "")
        # Use `description` field (new schema) — fallback to success-state for legacy
        desc = feat.get("description") or feat.get("success-state", "") or f"{PLACEHOLDER_MARK}: mô tả chức năng]"
        # Truncate at 200 chars (increased from 100) to preserve more context
        row[2].text = (desc[:200] + "...") if len(desc) > 200 else desc
        actors = feat.get("actors", [])
        row[3].text = ", ".join(actors) if actors else f"{PLACEHOLDER_MARK}: đối tượng sử dụng]"

    doc.add_paragraph()


def get_screenshots_for_step(screenshot_map, feat_id, step_no):
    """Return list of screenshots for specific feature + step."""
    feat_data = screenshot_map.get(feat_id, {})
    if isinstance(feat_data, list):
        screenshots = feat_data  # legacy flat format
    else:
        screenshots = feat_data.get("screenshots", [])
    return [sc for sc in screenshots if sc.get("step-no") == step_no]


def get_all_screenshots_for_feature(screenshot_map, feat_id):
    """Return all screenshots for a feature regardless of step-no."""
    feat_data = screenshot_map.get(feat_id, {})
    if isinstance(feat_data, list):
        return feat_data
    return feat_data.get("screenshots", [])


def build_noi_dung(doc, services, screenshot_map, screenshots_dir, all_error_cases):
    """Build Section II — Giới thiệu + Catalog + HDSD chi tiết."""
    all_features = [f for svc in services for f in svc.get("features", [])]

    # 1. Giới thiệu chung
    add_para(doc, "Giới thiệu chung", "A_Heading 2")
    add_para(doc, "Tổng quan hệ thống", "A_Heading 3")
    svc_names = ", ".join(s["display-name"] for s in services)
    add_para(
        doc,
        f"Tài liệu này hướng dẫn sử dụng các chức năng của hệ thống {svc_names}. "
        f"Hệ thống bao gồm {len(all_features)} chức năng chính được mô tả chi tiết dưới đây, "
        f"kèm theo hình ảnh minh họa và các bước thực hiện cụ thể.",
        CONTENT_STYLE,
    )
    doc.add_paragraph()

    add_para(doc, "Các nội dung khác", "A_Heading 3")
    add_para(
        doc,
        "Các quy tắc hiển thị của hệ thống:\n"
        "- Các trường bắt buộc được đánh dấu bằng ký hiệu * màu đỏ\n"
        "- Thông báo lỗi hiển thị bằng màu đỏ, thông báo thành công màu xanh\n"
        "- Định dạng ngày: DD/MM/YYYY\n\n"
        "Yêu cầu tối thiểu đối với người sử dụng:\n"
        "- Biết sử dụng máy vi tính và có kiến thức cơ bản về tin học\n"
        "- Trình duyệt: Chrome, Firefox, Edge phiên bản mới nhất",
        CONTENT_STYLE,
    )
    doc.add_paragraph()

    # 2. Giới thiệu các chức năng (catalog)
    add_para(doc, "Giới thiệu các chức năng", "A_Heading 2")
    for svc in services:
        add_para(doc, f"Các chức năng trong {svc['display-name']}", "A_Heading 3")
        add_feature_table(doc, svc.get("features", []))

    # 3. Hướng dẫn sử dụng
    add_para(doc, "Hướng dẫn sử dụng các chức năng", "A_Heading 2")
    for svc in services:
        add_para(doc, svc["display-name"], "A_Heading 3")

        for feat in svc.get("features", []):
            feat_id = feat.get("id", "")
            add_para(doc, feat.get("name", ""), "A_Heading 4")

            # Description first
            if feat.get("description"):
                add_para(doc, feat["description"], CONTENT_STYLE)

            # Preconditions
            if feat.get("preconditions"):
                add_para(doc, f"Điều kiện tiên quyết: {feat['preconditions']}", CONTENT_STYLE)

            steps = normalize_steps(feat)
            all_feat_screenshots = get_all_screenshots_for_feature(screenshot_map, feat_id)
            # Track which screenshots have been consumed by per-step embedding
            consumed_files = set()

            # Per-step: action → screenshots for that step → expected
            for step in steps:
                step_no = step.get("no", 0)
                add_para(doc, f"Bước {step_no}: {step.get('action', '')}", CONTENT_STYLE, bold=True)

                # Screenshots for this specific step (v2.0 schema)
                step_screenshots = get_screenshots_for_step(screenshot_map, feat_id, step_no)
                for sc in step_screenshots:
                    file = sc.get("file", "")
                    add_image(doc, screenshots_dir, file,
                              sc.get("description", f"Minh họa bước {step_no}"))
                    consumed_files.add(file)

                # Expected result
                if step.get("expected"):
                    add_para(doc, f"→ Kết quả: {step['expected']}", CONTENT_STYLE)

            # Leftover screenshots — step-no is None OR step-no does not match any step
            # (covers MIXED case where some sc have step-no, some don't)
            leftover = [
                sc for sc in all_feat_screenshots
                if sc.get("file", "") not in consumed_files
            ]
            if leftover:
                add_para(doc, "Hình ảnh minh họa bổ sung:", CONTENT_STYLE, bold=True)
                for sc in leftover:
                    add_image(doc, screenshots_dir, sc.get("file", ""),
                              sc.get("description", ""))

            # Success state summary
            if feat.get("success-state"):
                add_para(doc, f"Kết quả cuối cùng: {feat['success-state']}", CONTENT_STYLE)

            # Collect error cases for Section 4
            for err in normalize_error_cases(feat):
                all_error_cases.append({
                    "feature-name": feat.get("name", ""),
                    "feature-id": feat_id,
                    "condition": err.get("condition", "") if isinstance(err, dict) else str(err),
                    "message": err.get("message", "") if isinstance(err, dict) else "",
                })

            doc.add_paragraph()


def build_troubleshooting(doc, all_error_cases):
    """Build Section 4 — Các vấn đề thường gặp khi sử dụng (from error-cases)."""
    add_para(doc, "Các vấn đề thường gặp khi sử dụng", "A_Heading 2")

    if not all_error_cases:
        add_placeholder_para(doc, "Chưa phát hiện lỗi thường gặp — bổ sung sau khi triển khai thực tế")
        return

    add_para(
        doc,
        "Dưới đây là danh sách các vấn đề có thể phát sinh khi sử dụng hệ thống và cách xử lý:",
        CONTENT_STYLE,
    )

    # Deduplicate by (condition, message) — same error appearing in N features
    # becomes 1 row with aggregated feature list (Issue #12 fix)
    deduped = {}
    for err in all_error_cases:
        key = (err["condition"].strip().lower(), err["message"].strip().lower())
        if key not in deduped:
            deduped[key] = {
                "condition": err["condition"],
                "message": err["message"],
                "features": [],
            }
        if err["feature-name"] not in deduped[key]["features"]:
            deduped[key]["features"].append(err["feature-name"])

    # Sort by number of features affected (most common first) then alphabetical
    sorted_errs = sorted(deduped.values(),
                         key=lambda x: (-len(x["features"]), x["condition"]))

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    _set_table_full_width(table)
    hdr = table.rows[0].cells
    for i, title in enumerate(["STT", "Chức năng liên quan", "Tình huống lỗi", "Cách xử lý"]):
        hdr[i].text = title
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True

    for stt, err in enumerate(sorted_errs, start=1):
        row = table.add_row().cells
        row[0].text = str(stt)
        # Show up to 3 features; if more, show "X features (including A, B, C)"
        feats = err["features"]
        if len(feats) <= 3:
            row[1].text = ", ".join(feats)
        else:
            row[1].text = f"{len(feats)} chức năng (trong đó: {', '.join(feats[:3])}...)"
        row[2].text = err["condition"]
        row[3].text = err["message"] or f"{PLACEHOLDER_MARK}: hướng dẫn xử lý]"

    if len(all_error_cases) > len(sorted_errs):
        doc.add_paragraph()
        add_para(
            doc,
            f"(Đã gộp {len(all_error_cases)} error-cases thành {len(sorted_errs)} "
            f"tình huống duy nhất — các lỗi giống nhau xuất hiện ở nhiều chức năng được aggregate.)",
            CONTENT_STYLE,
        )

    doc.add_paragraph()
    add_para(
        doc,
        f"Liên hệ hỗ trợ kỹ thuật: {PLACEHOLDER_MARK}: email/hotline đơn vị phát triển]",
        CONTENT_STYLE,
    )


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True)
    parser.add_argument("--flow-report", required=True)
    parser.add_argument("--stack-report", help="Optional: intel/stack-report.json for richer Section I")
    parser.add_argument("--screenshot-map", required=True)
    parser.add_argument("--screenshots-dir", required=True)
    parser.add_argument("--output-dir", required=True)
    parser.add_argument("--project-name", required=True)
    parser.add_argument("--client-name", default="", help="Tên khách hàng (có thể khác dev-unit)")
    parser.add_argument("--dev-unit", required=True, help="Đơn vị phát triển (công ty xây dựng phần mềm)")
    args = parser.parse_args()

    today = datetime.now().strftime("%d/%m/%Y")

    for path, label in [(args.template, "template"), (args.flow_report, "flow-report")]:
        if not os.path.exists(path):
            print(f"ERROR: {label} not found: {path}")
            return 1

    with open(args.flow_report, encoding="utf-8") as f:
        flow = json.load(f)

    # Load stack report if provided (for Section I enrichment)
    stack = {}
    if args.stack_report and os.path.exists(args.stack_report):
        with open(args.stack_report, encoding="utf-8") as f:
            stack = json.load(f)

    screenshot_map = load_screenshot_map(args.screenshot_map)
    os.makedirs(args.output_dir, exist_ok=True)

    # Resolve services (support split-mode)
    if flow.get("split-mode"):
        print("WARNING: flow-report is split-mode — loading parts...")
        parts_dir = os.path.dirname(args.flow_report)
        services_by_name = {}
        for part in flow.get("parts", []):
            part_path = os.path.join(parts_dir, part["file"])
            if os.path.exists(part_path):
                with open(part_path, encoding="utf-8") as f:
                    part_data = json.load(f)
                for svc in part_data.get("services", []):
                    existing = services_by_name.get(svc["name"])
                    if existing:
                        existing["features"].extend(svc.get("features", []))
                    else:
                        services_by_name[svc["name"]] = svc
        raw_services = list(services_by_name.values())
    else:
        raw_services = flow.get("services", [])

    if not raw_services:
        raw_services = [{
            "name": "he-thong",
            "display-name": args.project_name,
            "features": flow.get("features", [])
        }]

    services = [normalize_service(s, args.project_name) for s in raw_services]

    # Scalability warning for large codebases
    total_features = sum(len(s.get("features", [])) for s in services)
    if total_features > 50:
        print(f"⚠️ LARGE CODEBASE: {total_features} features — generated docx may exceed 100 pages.")
        print("   Consider generating docs per module to reduce review burden.")

    client_name = args.client_name or args.project_name

    # Collect per-service metrics for final JSON report (Issue #5 + #9 fixes)
    per_service_results = []

    # Generate one docx per service (per doc-manual-writer logic)
    for service in services:
        svc_slug = service["name"].lower().replace(" ", "-").replace("_", "-")
        svc_features = service.get("features", [])

        if not svc_features:
            print(f"SKIP {svc_slug}: no features")
            per_service_results.append({
                "service": svc_slug,
                "status": "skipped",
                "reason": "no features in flow-report",
                "file": None,
                "metrics": {},
            })
            continue

        try:
            doc = Document(args.template)

            # Step 1: Replace all placeholders (cover + header + footer)
            replace_placeholders(doc, args.project_name, client_name, args.dev_unit, today)

            # Step 2: Fill revision history + signing pages
            fill_revision_history(doc, today)
            fill_signing_pages(doc, args.dev_unit, client_name)

            # Step 3: Fill Section I (Tổng quan tài liệu)
            single_service_flow = {"services": [service]}
            fill_section_i(doc, args.project_name, single_service_flow, stack)

            # Step 4: Clear after "NỘI DUNG" heading
            cleared = clear_after_heading(doc, "NỘI DUNG")
            if not cleared:
                print(f"WARNING [{svc_slug}]: 'NỘI DUNG' heading not found — appending at end")

            # Step 5: Rebuild Section II (Giới thiệu + Catalog + HDSD) + collect errors
            all_errors = []
            build_noi_dung(doc, [service], screenshot_map, args.screenshots_dir, all_errors)

            # Step 6: Rebuild Section 4 (Các vấn đề thường gặp) — with dedup (Issue #12)
            build_troubleshooting(doc, all_errors)

            # Step 7: Mark TOC dirty for Word auto-update
            mark_toc_dirty(doc)

            out_path = os.path.join(args.output_dir, f"huong-dan-su-dung-{svc_slug}.docx")
            doc.save(out_path)

            # Calculate depth metrics (Issue #5 fix)
            metrics = _calculate_doc_metrics(doc, len(svc_features))
            size_mb = os.path.getsize(out_path) / (1024 * 1024)

            # Count CẦN BỔ SUNG placeholders
            cabosung_count = _count_placeholders_in_doc(doc)

            print(f"Saved: {out_path} ({size_mb:.1f} MB, {len(svc_features)} features, "
                  f"~{metrics['estimated_pages']:.0f} pages, "
                  f"{metrics['pages_per_feature']:.1f} pages/feature avg)")

            if metrics['pages_per_feature'] < 2.5:
                print(f"   ⚠️ SPARSE OUTPUT [{svc_slug}]: avg {metrics['pages_per_feature']:.1f} pages/feature "
                      f"(target ≥3). Writers may have produced sparse content.")

            if size_mb > 20:
                print(f"   ⚠️ Large file [{svc_slug}] — may be slow to open in Word")

            per_service_results.append({
                "service": svc_slug,
                "status": "ok" if metrics['pages_per_feature'] >= 2.5 else "sparse",
                "file": out_path,
                "metrics": {
                    "size_mb": round(size_mb, 2),
                    "features": len(svc_features),
                    "word_count": metrics['word_count'],
                    "paragraph_count": metrics['paragraph_count'],
                    "table_count": metrics['table_count'],
                    "image_count": metrics['image_count'],
                    "estimated_pages": round(metrics['estimated_pages'], 1),
                    "pages_per_feature": round(metrics['pages_per_feature'], 2),
                    "cabosung_markers": cabosung_count,
                },
            })

        except Exception as e:
            print(f"ERROR [{svc_slug}]: {type(e).__name__}: {e}")
            per_service_results.append({
                "service": svc_slug,
                "status": "failed",
                "reason": f"{type(e).__name__}: {str(e)[:200]}",
                "file": None,
                "metrics": {},
            })

    # Write per-service JSON report for exporter (Issue #9 fix)
    report_path = os.path.join(args.output_dir, "manual-fill-report.json")
    with open(report_path, "w", encoding="utf-8") as f:
        json.dump({
            "generated_at": datetime.now().isoformat(),
            "total_services": len(services),
            "succeeded": sum(1 for r in per_service_results if r["status"] == "ok"),
            "sparse": sum(1 for r in per_service_results if r["status"] == "sparse"),
            "failed": sum(1 for r in per_service_results if r["status"] == "failed"),
            "skipped": sum(1 for r in per_service_results if r["status"] == "skipped"),
            "services": per_service_results,
        }, f, ensure_ascii=False, indent=2)
    print(f"\nReport: {report_path}")

    return 0


# ── Helper functions for metrics (Issue #5) ───────────────────────────────────

def _calculate_doc_metrics(doc, feature_count):
    """Estimate page count and detail depth from document structure."""
    word_count = 0
    paragraph_count = 0
    image_count = 0
    table_count = len(doc.tables)

    for para in doc.paragraphs:
        if para.text.strip():
            word_count += len(para.text.split())
            paragraph_count += 1

    # Count paragraphs in tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        word_count += len(para.text.split())

    # Count inline images (approximate)
    for para in doc.paragraphs:
        for run in para.runs:
            if run._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'):
                image_count += len(run._element.findall(
                    './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'
                ))

    # Page estimation heuristic:
    # - ~400 words per page (body text Times New Roman 13pt, 1.5 line spacing)
    # - Each image ≈ 0.5 page
    # - Each large table (>10 rows) ≈ 1 page
    large_tables = sum(1 for t in doc.tables if len(t.rows) > 10)
    estimated_pages = max(1.0, (word_count / 400.0) + (image_count * 0.5) + large_tables)
    pages_per_feature = estimated_pages / max(feature_count, 1)

    return {
        "word_count": word_count,
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "image_count": image_count,
        "estimated_pages": estimated_pages,
        "pages_per_feature": pages_per_feature,
    }


def _count_placeholders_in_doc(doc):
    """Count [CẦN BỔ SUNG markers across entire document."""
    count = 0
    for para in _iter_all_paragraphs(doc):
        count += para.text.count(PLACEHOLDER_MARK)
    return count


if __name__ == "__main__":
    sys.exit(main())
