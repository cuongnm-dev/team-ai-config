#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
render_docx.py — Jinja2-based .docx rendering via docxtpl.

REPLACES: fill_docx_engine.py (~900 lines of custom XML manipulation)

Philosophy: Templates are forks of ETC masters with {{ }} / {% %} tags added.
DocxTpl (Jinja2 for Word) does the actual rendering.

Template authoring convention:
  - {{ project.display_name }}        → simple substitution
  - {%tr for f in features %} ... {%tr endfor %}   → loop table rows
  - {%p if feat.preconditions %} ... {%p endif %}  → conditional paragraphs
  - {%p for step in steps %}{{ step.action }}{%p endfor %}   → loop paragraphs
  - Image: passed via context as InlineImage(tpl, path, width)

Post-processing after render:
  - TOC dirty flag (so Word refreshes on open)
  - Embed screenshots as InlineImage in context pre-render

Usage:
  python render_docx.py \
    --template       path/to/template-jinja.docx \
    --data           content-data.json \
    --output         out/filled.docx \
    [--screenshots-dir path/]   # optional, enables InlineImage for screenshots
    [--report        out/render-report.json]
"""
from __future__ import annotations

import argparse
import json
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

try:
    from docxtpl import DocxTemplate, InlineImage
    from docx.shared import Inches, Mm
except ImportError:
    print("ERROR: docxtpl required. Run: pip install docxtpl python-docx", file=sys.stderr)
    sys.exit(2)

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx required", file=sys.stderr)
    sys.exit(2)


WNS_QN = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# ─────────────────────────── Report ───────────────────────────

@dataclass
class RenderReport:
    template: str = ""
    output: str = ""
    screenshots_embedded: int = 0
    screenshots_missing: int = 0
    diagrams_embedded: int = 0
    diagrams_missing: int = 0
    tokens_substituted: int = 0    # approximation — count {{ }} tokens
    warnings: list[str] = field(default_factory=list)
    errors: list[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        d = {k: v for k, v in self.__dict__.items()}
        d["status"] = "ok" if not self.errors else "failed"
        return d


# ─────────────────────────── Image context pre-processor ───────────────────────────

def build_diagram_context(tpl: DocxTemplate, data: dict, diagrams_dir: Path | None,
                           report: RenderReport) -> dict:
    """Resolve data['diagrams'][<key>] strings → `diagram_<key>` InlineImage refs.

    Template uses `{{ diagram_component }}`, `{{ diagram_erd }}`, etc. to embed
    rendered Mermaid PNGs. Source strings in data['diagrams'] are Mermaid code
    (kept for re-render); the PNGs must be pre-rendered by render_mermaid.py.
    """
    diagrams = data.get("diagrams") or {}
    if not diagrams:
        # No diagram block — inject None so template conditionals skip cleanly
        for key in ("component", "erd", "deployment", "sequence", "integration",
                    "business", "security"):
            data[f"diagram_{key}"] = None
        return data

    for key, src in diagrams.items():
        field = f"diagram_{key}"
        png = (diagrams_dir / f"{key}.png") if diagrams_dir else None
        if png and png.exists():
            data[field] = InlineImage(tpl, str(png), width=Inches(6.0))
            if hasattr(report, "diagrams_embedded"):
                report.diagrams_embedded += 1
        else:
            data[field] = None
            if hasattr(report, "diagrams_missing"):
                report.diagrams_missing += 1
    return data


def build_image_context(tpl: DocxTemplate, data: dict, screenshots_dir: Path | None,
                         report: RenderReport) -> dict:
    """Walk services[].features[].steps[] — replace screenshot filenames
    with InlineImage objects (so Jinja renders them as embedded images).

    If file missing, replace with None — template should have
    {%p if step.screenshot %}...{%p endif %} guard to skip missing images.
    """
    if not screenshots_dir or not screenshots_dir.exists():
        # Mark all screenshots as missing but still replace strings with None
        # so Jinja conditionals can skip cleanly
        for svc in data.get("services", []):
            for feat in svc.get("features", []):
                for step in feat.get("steps", []):
                    if step.get("screenshot"):
                        report.screenshots_missing += 1
                    step["screenshot_image"] = None
        return data

    for svc in data.get("services", []):
        for feat in svc.get("features", []):
            for step in feat.get("steps", []):
                fn = step.get("screenshot")
                if not fn:
                    step["screenshot_image"] = None
                    continue
                candidates = [screenshots_dir / fn]
                # Try alternate extensions (JPEG from post-processor)
                stem = Path(fn).stem
                for ext in (".png", ".jpg", ".jpeg", ".webp"):
                    candidates.append(screenshots_dir / (stem + ext))
                resolved = next((c for c in candidates if c.exists()), None)
                if resolved:
                    step["screenshot_image"] = InlineImage(
                        tpl, str(resolved), width=Inches(5.5)
                    )
                    report.screenshots_embedded += 1
                else:
                    step["screenshot_image"] = None
                    report.screenshots_missing += 1
    return data


# ─────────────────────────── Post-process ───────────────────────────

def mark_toc_dirty(docx_path: Path):
    """Set updateFields=true + mark TOC fields as dirty so Word auto-refreshes."""
    doc = Document(docx_path)
    settings = doc.settings.element
    upd = settings.find(f"{WNS_QN}updateFields")
    if upd is None:
        upd = OxmlElement("w:updateFields")
        settings.append(upd)
    upd.set(qn("w:val"), "true")
    for fc in doc.element.body.iter(f"{WNS_QN}fldChar"):
        if fc.get(qn("w:fldCharType")) == "begin":
            fc.set(qn("w:dirty"), "true")
    doc.save(docx_path)


def strip_orphan_media(docx_path: Path) -> int:
    """Remove media files no longer referenced by document.xml.

    Handles the case of forked templates carrying orphan images (e.g. TKCS
    inherited from v1.2) — after DocxTpl render clears original body,
    those media become orphan and can be safely stripped.
    """
    import zipfile
    import re
    import shutil

    REL_FILE = "word/_rels/document.xml.rels"
    DOC_FILE = "word/document.xml"

    try:
        with zipfile.ZipFile(docx_path, "r") as zin:
            doc_xml = zin.read(DOC_FILE).decode("utf-8")
            rels_xml = zin.read(REL_FILE).decode("utf-8")
    except (KeyError, zipfile.BadZipFile):
        return 0

    used_rids = set(re.findall(r'r:(?:embed|link|id)="([^"]+)"', doc_xml))
    rel_pattern = re.compile(
        r'<Relationship\s+[^/]*Id="([^"]+)"[^/]*Target="(media/[^"]+|embeddings/[^"]+)"[^/]*/>',
        re.DOTALL,
    )
    orphan_rels = {}
    for m in rel_pattern.finditer(rels_xml):
        rid, target = m.group(1), m.group(2)
        if rid not in used_rids:
            orphan_rels[rid] = target
    if not orphan_rels:
        return 0

    new_rels = rels_xml
    for rid in orphan_rels:
        new_rels = re.sub(
            rf'<Relationship\s+[^/]*Id="{re.escape(rid)}"[^/]*/>',
            "", new_rels,
        )
    orphan_targets = {f"word/{t}" for t in orphan_rels.values()}

    tmp = docx_path.with_suffix(".tmp.docx")
    removed = 0
    with zipfile.ZipFile(docx_path, "r") as zin:
        with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename in orphan_targets:
                    removed += 1
                    continue
                if item.filename == REL_FILE:
                    zout.writestr(item, new_rels.encode("utf-8"))
                    continue
                zout.writestr(item, zin.read(item.filename))
    shutil.move(str(tmp), str(docx_path))
    return removed


# ─────────────────────────── Main render ───────────────────────────

def render(template_path: Path, data_path: Path, output_path: Path,
           screenshots_dir: Path | None = None,
           diagrams_dir: Path | None = None) -> RenderReport:
    report = RenderReport(template=str(template_path), output=str(output_path))

    if not template_path.exists():
        report.errors.append(f"Template not found: {template_path}")
        return report

    # Load data
    try:
        data = json.loads(data_path.read_text(encoding="utf-8"))
    except Exception as e:
        report.errors.append(f"Cannot parse data: {e}")
        return report

    # Add `meta.today` fallback if missing
    data.setdefault("meta", {}).setdefault("today", "")

    # Pre-compute `all_features` flat list (features with service_name embedded)
    # — simplifies templates by avoiding nested {%p for service %} + {%tr for feat %}
    # which docxtpl/Jinja cannot parse reliably.
    all_features = []
    for svc in data.get("services", []):
        svc_name = svc.get("display_name", "")
        for feat in svc.get("features", []):
            flat = dict(feat)
            flat["service_name"] = svc_name
            all_features.append(flat)
    data["all_features"] = all_features

    # Open template
    try:
        tpl = DocxTemplate(str(template_path))
    except Exception as e:
        report.errors.append(f"Cannot open template: {e}")
        return report

    # Pre-process screenshots + diagrams into InlineImage objects
    data = build_image_context(tpl, data, screenshots_dir, report)
    data = build_diagram_context(tpl, data, diagrams_dir, report)

    # Render
    try:
        tpl.render(data)
    except Exception as e:
        report.errors.append(f"Render failed: {e}")
        return report

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        tpl.save(str(output_path))
    except Exception as e:
        report.errors.append(f"Save failed: {e}")
        return report

    # Post-process
    try:
        mark_toc_dirty(output_path)
    except Exception as e:
        report.warnings.append(f"TOC dirty failed: {e}")

    try:
        orphans = strip_orphan_media(output_path)
        if orphans > 0:
            report.warnings.append(f"Stripped {orphans} orphan media files")
    except Exception as e:
        report.warnings.append(f"Orphan cleanup failed: {e}")

    # Validate output
    try:
        doc = Document(output_path)
        # Check for residual Jinja tokens
        full_text = "\n".join(p.text for p in doc.paragraphs)
        residuals = []
        for marker in ("{{", "}}", "{%", "%}"):
            if marker in full_text:
                residuals.append(marker)
        if residuals:
            report.warnings.append(f"Residual Jinja markers in output: {residuals}")
    except Exception as e:
        report.warnings.append(f"Output validation failed: {e}")

    return report


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", required=True)
    ap.add_argument("--data", required=True)
    ap.add_argument("--output", required=True)
    ap.add_argument("--screenshots-dir", default=None)
    ap.add_argument("--diagrams-dir", default=None)
    ap.add_argument("--report", default=None)
    args = ap.parse_args()

    report = render(
        Path(args.template),
        Path(args.data),
        Path(args.output),
        Path(args.screenshots_dir) if args.screenshots_dir else None,
        Path(args.diagrams_dir) if args.diagrams_dir else None,
    )

    d = report.to_dict()
    if args.report:
        Path(args.report).parent.mkdir(parents=True, exist_ok=True)
        Path(args.report).write_text(
            json.dumps(d, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

    print(f"Template: {report.template}")
    print(f"Output:   {report.output}")
    print(f"Screenshots: {report.screenshots_embedded} embedded, "
          f"{report.screenshots_missing} missing")
    if report.warnings:
        print(f"Warnings: {len(report.warnings)}")
        for w in report.warnings:
            print(f"  - {w}")
    if report.errors:
        print(f"ERRORS: {len(report.errors)}")
        for e in report.errors:
            print(f"  ✗ {e}")
        sys.exit(1)
    print("OK")


if __name__ == "__main__":
    main()
