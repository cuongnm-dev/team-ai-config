#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
jinjafy_templates.py — Fork ETC templates by adding Jinja2 (docxtpl) tags.

Takes original ETC template, adds:
  - {{ variable }} tags for placeholder replacements
  - {%tr for %} / {%tr endfor %} in table rows for dynamic tables
  - {%p for %} / {%p endfor %} in paragraphs for content loops
  - {% if %} / {% endif %} for conditionals

Outputs to templates/ with the standard filename (replaces pre-Jinja version).

IMPORTANT: this script is the ONLY way templates get Jinja tags added.
Don't hand-edit the .docx — re-run this script if template structure changes.

Usage:
  python jinjafy_templates.py hdsd  --source path/to/huong-dan-su-dung.docx
  python jinjafy_templates.py tkkt  --source templates/huong-dan-su-dung.docx   # uses HDSD as base + outline YAML
  python jinjafy_templates.py tkcs  --source templates/huong-dan-su-dung.docx   # uses HDSD as base + outline YAML

TKKT + TKCS body content is driven by outline+fieldmap YAML files in schemas/:
  - schemas/tkkt-outline.yaml + schemas/tkkt-fieldmap.yaml
  - schemas/tkcs-outline.yaml + schemas/tkcs-fieldmap.yaml
"""
from __future__ import annotations
import argparse
import re
import shutil
import sys
from pathlib import Path
from typing import Callable

import yaml
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

SCHEMA_DIR = Path(__file__).parent.parent / "schemas"

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WNS_QN = f"{{{WNS}}}"


# ─────────────────────────── Helpers ───────────────────────────

def para_text(p_elem) -> str:
    return "".join(t.text or "" for t in p_elem.findall(f".//{WNS_QN}t"))


def para_style(p_elem) -> str:
    ps = p_elem.find(f".//{WNS_QN}pStyle")
    return ps.get(f"{WNS_QN}val", "") if ps is not None else ""


def set_para_text(p_elem, new_text: str):
    """Rebuild paragraph runs with single run = new_text, keep first run formatting."""
    runs = p_elem.findall(f".//{WNS_QN}r")
    if runs:
        first_run = runs[0]
        for t in first_run.findall(f"{WNS_QN}t"):
            first_run.remove(t)
        new_t = OxmlElement("w:t")
        new_t.text = new_text
        new_t.set(qn("xml:space"), "preserve")
        first_run.append(new_t)
        for r in runs[1:]:
            for t in r.findall(f"{WNS_QN}t"):
                t.text = ""
    else:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = new_text
        t.set(qn("xml:space"), "preserve")
        r.append(t)
        p_elem.append(r)


def walk_all_paragraphs(doc):
    """Yield every paragraph including headers, footers, table cells."""
    yield from doc.paragraphs
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nested in cell.tables:
                    for nrow in nested.rows:
                        for ncell in nrow.cells:
                            yield from ncell.paragraphs
    for section in doc.sections:
        for part in (section.header, section.first_page_header, section.even_page_header,
                     section.footer, section.first_page_footer, section.even_page_footer):
            if part is None:
                continue
            yield from part.paragraphs
            for tbl in part.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        yield from cell.paragraphs


def replace_everywhere(doc, replacements: list[tuple[str, str]]) -> int:
    """Walk all paragraphs, apply split-run-safe string replacements."""
    count = 0
    for p in walk_all_paragraphs(doc):
        full = "".join(r.text or "" for r in p.runs)
        if not full:
            continue
        new = full
        for old, val in replacements:
            if old in new:
                new = new.replace(old, val)
        if new != full and p.runs:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ""
            count += 1
    return count


def add_paragraph_after(target_elem, text: str, style: str | None = None,
                        bold: bool = False, italic: bool = False):
    """Insert new paragraph right after target_elem in XML tree."""
    new_p = OxmlElement("w:p")
    if style:
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), style)
        pPr.append(pStyle)
        new_p.append(pPr)
    new_r = OxmlElement("w:r")
    if bold or italic:
        new_rPr = OxmlElement("w:rPr")
        if bold:
            new_rPr.append(OxmlElement("w:b"))
        if italic:
            new_rPr.append(OxmlElement("w:i"))
        new_r.append(new_rPr)
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_t.set(qn("xml:space"), "preserve")
    new_r.append(new_t)
    new_p.append(new_r)
    target_elem.addnext(new_p)
    return new_p


def clear_after_heading(doc, heading_text: str) -> int:
    """Delete all body elements after the heading matching text."""
    body = doc.element.body
    children = list(body)
    needle = heading_text.upper()
    marker_idx = None
    for i, ch in enumerate(children):
        tag = ch.tag.split("}")[-1] if "}" in ch.tag else ch.tag
        if tag != "p":
            continue
        style = para_style(ch).upper()
        text = para_text(ch).upper()
        if "HEADING" in style and needle in text:
            marker_idx = i
            break
    if marker_idx is None:
        return 0
    removed = 0
    for ch in children[marker_idx + 1:]:
        tag = ch.tag.split("}")[-1] if "}" in ch.tag else ch.tag
        if tag == "sectPr":
            continue
        body.remove(ch)
        removed += 1
    return removed


# ─────────────────────────── Template 1: HDSD ───────────────────────────

def jinjafy_hdsd(source: Path, dest: Path, today: str = "{{ meta.today }}"):
    """Fork ETC huong-dan-su-dung.docx with Jinja2 tags."""
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source, dest)
    doc = Document(dest)

    # 1. Cover + header/footer placeholders
    replacements = [
        ("<Tên dự án>", "{{ project.display_name }}"),
        ("<Mã dự án>", "{{ project.code }}"),
        ("Phiên bản: <xx>", "Phiên bản: {{ meta.version or '1.0' }}"),
        ("HỆ THỐNG QUẢN LÝ", "{{ project.display_name|upper }}"),
        ("CÔNG TY CỔ PHẦN HỆ THỐNG CÔNG NGHỆ ETC", "{{ dev_unit|upper }}"),
        ("<Họ tên>", "[CẦN BỔ SUNG: Tên người ký]"),
        ("<Chức danh>", "[CẦN BỔ SUNG: Chức danh]"),
        ("<Màn hình 1>", "[Minh họa màn hình]"),
        ("Ban hành: dd/mm/yyyy", "Ban hành: {{ meta.today }}"),
    ]
    n = replace_everywhere(doc, replacements)
    print(f"  Cover placeholders: {n} paragraphs updated")

    # 2. Change history (T[1] row 1) — fill with Jinja
    if len(doc.tables) > 1:
        t = doc.tables[1]
        if len(t.rows) >= 2:
            row = t.rows[1]
            cells = row.cells
            if len(cells) >= 5:
                cells[0].text = "{{ meta.today }}"
                cells[1].text = "Toàn bộ tài liệu"
                cells[2].text = "A"
                cells[3].text = "Tạo mới tài liệu"
                cells[4].text = "{{ meta.version or '1.0' }}"
                print(f"  T[1] change history: row 1 Jinja-fied")

    # 3. Abbreviations table T[4] — replace rows with Jinja loop
    if len(doc.tables) > 4:
        jinjafy_loop_table(doc.tables[4], "overview.terms",
                           ["short", "full", "explanation"])
        print(f"  T[4] abbreviations: loop added")

    # 4. Related docs T[5] — same pattern
    if len(doc.tables) > 5:
        jinjafy_loop_table(doc.tables[5], "overview.references",
                           ["stt", "name", "ref"])
        print(f"  T[5] related docs: loop added")

    # 5. Section I fills (Mục đích tài liệu, Phạm vi tài liệu)
    fill_section_i_after_heading(doc, "Mục đích tài liệu",
                                 "{{ overview.purpose }}")
    fill_section_i_after_heading(doc, "Phạm vi tài liệu",
                                 "{{ overview.scope }}")
    print(f"  Section I: purpose + scope replaced with Jinja expressions")

    # 6. Section II — clear after NỘI DUNG heading, append Jinja layout
    removed = clear_after_heading(doc, "NỘI DUNG")
    print(f"  Cleared {removed} elements after NỘI DUNG")
    append_hdsd_section_ii(doc)
    print(f"  Appended Section II Jinja layout (~50 paragraphs)")

    doc.save(dest)
    size_kb = dest.stat().st_size // 1024
    print(f"  Saved {dest} ({size_kb} KB)")


def jinjafy_loop_table(table, loop_var: str, columns: list[str]):
    """Wrap table data rows in {%tr for item %} / {%tr endfor %} — 3-row pattern.

    docxtpl requires {%tr for %} and {%tr endfor %} to be ALONE in their own rows
    (separate from data rows). Structure after this function:
      Row 0: header (untouched)
      Row 1: {%tr for item in loop_var %} alone in cell 0, other cells empty — REMOVED at render
      Row 2: data row with {{ }} expressions — REPEATED for each item
      Row 3: {%tr endfor %} alone in cell 0 — REMOVED at render
    """
    # Delete all data rows except header
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    singular = loop_var.split(".")[-1].rstrip("s")  # "terms" → "term"
    n_cols = len(table.columns)

    # Row 1: for directive alone
    for_row = table.add_row()
    for_row.cells[0].text = f"{{%tr for {singular} in {loop_var} %}}"
    for c in for_row.cells[1:]:
        c.text = ""

    # Row 2: data row with {{ }} substitutions
    data_row = table.add_row()
    for i in range(min(n_cols, len(columns))):
        col_key = columns[i]
        if col_key == "@index":
            data_row.cells[i].text = "{{ loop.index }}"
        else:
            data_row.cells[i].text = f"{{{{ {singular}.{col_key} }}}}"

    # Row 3: endfor directive alone
    end_row = table.add_row()
    end_row.cells[0].text = "{%tr endfor %}"
    for c in end_row.cells[1:]:
        c.text = ""


def fill_section_i_after_heading(doc, heading_text: str, jinja_expr: str):
    """Find heading, replace content paragraphs between it and next heading
    with a single Jinja expression paragraph."""
    body = doc.element.body
    children = list(body)
    needle = heading_text.upper()

    # Find heading
    heading_idx = None
    for i, ch in enumerate(children):
        tag = ch.tag.split("}")[-1] if "}" in ch.tag else ch.tag
        if tag != "p":
            continue
        if "HEADING" in para_style(ch).upper() and needle in para_text(ch).upper():
            heading_idx = i
            break
    if heading_idx is None:
        return False

    # Find next heading (end boundary)
    end_idx = len(children)
    for j in range(heading_idx + 1, len(children)):
        ch = children[j]
        tag = ch.tag.split("}")[-1] if "}" in ch.tag else ch.tag
        if tag == "p" and "HEADING" in para_style(ch).upper():
            end_idx = j
            break

    # Delete everything between heading and next heading
    for ch in children[heading_idx + 1:end_idx]:
        tag = ch.tag.split("}")[-1] if "}" in ch.tag else ch.tag
        if tag in ("p", "tbl"):
            body.remove(ch)

    # Insert single Jinja paragraph with ETC_Content style
    add_paragraph_after(children[heading_idx], jinja_expr, style="ETC_Content")
    return True


def append_hdsd_section_ii(doc):
    """Append Jinja2 template block for Section II content after NỘI DUNG.

    Design choice: avoid nested {%p for %} + {%tr for %} — docxtpl/Jinja
    can't parse that mix. Instead use FLAT tables (no outer p-loop wrapping
    a table) and pure paragraph nesting for user manual.

    Requires render_docx.py to pre-compute `all_features` from services[].
    """
    def H(text, level):
        style = f"A_Heading {level}" if level > 1 else "A_HEADING 1"
        return doc.add_paragraph(text, style=style)

    def P(text, style="ETC_Content", bold=False, italic=False):
        p = doc.add_paragraph(style=style)
        run = p.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        return p

    # ─ 2.1 Giới thiệu chung ─
    H("Giới thiệu chung", 2)
    H("Tổng quan chương trình", 3)
    P("{{ architecture.system_overview or overview.system_description }}")
    H("Các nội dung khác", 3)
    P("{{ overview.conventions }}")

    # ─ 2.2 Giới thiệu các chức năng (FLAT combined catalog) ─
    H("Giới thiệu các chức năng", 2)
    H("Danh mục chức năng toàn hệ thống", 3)
    tbl = doc.add_table(rows=4, cols=5)
    try:
        tbl.style = "Table Grid"
    except KeyError:
        pass
    hdr = tbl.rows[0].cells
    for i, label in enumerate(["STT", "Phân hệ", "Chức năng", "Mô tả", "Đối tượng"]):
        hdr[i].text = label
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
    # Row 1: for-directive alone
    tbl.rows[1].cells[0].text = "{%tr for feat in all_features %}"
    # Row 2: data row — repeated per feature
    data = tbl.rows[2].cells
    data[0].text = "{{ loop.index }}"
    data[1].text = "{{ feat.service_name }}"
    data[2].text = "{{ feat.name }}"
    data[3].text = "{{ feat.description }}"
    data[4].text = "{{ feat.actors|join(', ') }}"
    # Row 3: endfor alone
    tbl.rows[3].cells[0].text = "{%tr endfor %}"

    # ─ 2.3 Hướng dẫn sử dụng các chức năng (PARAGRAPH-ONLY nesting) ─
    H("Hướng dẫn sử dụng các chức năng hệ thống", 2)
    P("{%p for service in services %}")
    H("{{ service.display_name }}", 3)
    P("{%p for feat in service.features %}")
    H("{{ feat.name }}", 4)
    P("{{ feat.description }}")
    P("{%p if feat.preconditions %}")
    P("Điều kiện tiên quyết: {{ feat.preconditions }}", italic=True)
    P("{%p endif %}")

    # UI elements as BULLETED LIST (paragraphs, not table — avoids nesting clash)
    P("{%p if feat.ui_elements %}")
    P("Các thành phần trên màn hình:", bold=True)
    P("{%p for elem in feat.ui_elements %}")
    P("• {{ elem.label }} ({{ elem.type }}) — {{ elem.rules or '' }}")
    P("{%p endfor %}")
    P("{%p endif %}")

    # Steps as paragraphs (each step: action + image + expected)
    P("{%p for step in feat.steps %}")
    P("Bước {{ step.no }}: {{ step.action }}", bold=True)
    P("{%p if step.screenshot_image %}")
    P("{{ step.screenshot_image }}")
    P("{%p endif %}")
    P("{%p if step.expected %}")
    P("→ Kết quả: {{ step.expected }}")
    P("{%p endif %}")
    P("{%p endfor %}")

    # Error cases (bulleted paragraphs — avoid table in nest)
    P("{%p if feat.error_cases %}")
    P("Các trường hợp lỗi:", bold=True)
    P("{%p for err in feat.error_cases %}")
    P("• Bước {{ err.trigger_step }}: {{ err.condition }} → {{ err.message }}")
    P("{%p endfor %}")
    P("{%p endif %}")

    P("{%p endfor %}")  # end features
    P("{%p endfor %}")  # end services

    # ─ 2.4 Các vấn đề thường gặp (FLAT table, 3-row pattern) ─
    H("Các vấn đề thường gặp khi sử dụng", 2)
    tbl3 = doc.add_table(rows=4, cols=4)
    try:
        tbl3.style = "Table Grid"
    except KeyError:
        pass
    hdr3 = tbl3.rows[0].cells
    for i, label in enumerate(["STT", "Tình huống", "Nguyên nhân", "Cách xử lý"]):
        hdr3[i].text = label
        for r in hdr3[i].paragraphs[0].runs:
            r.bold = True
    tbl3.rows[1].cells[0].text = "{%tr for item in troubleshooting %}"
    dr = tbl3.rows[2].cells
    dr[0].text = "{{ loop.index }}"
    dr[1].text = "{{ item.situation }}"
    dr[2].text = "{{ item.cause }}"
    dr[3].text = "{{ item.resolution }}"
    tbl3.rows[3].cells[0].text = "{%tr endfor %}"


# ─────────────────────────── Outline-driven body builder ───────────────────────────

def _load_yaml(path: Path) -> dict:
    with path.open(encoding="utf-8") as f:
        return yaml.safe_load(f)


def _heading_style(level: int) -> str:
    """Map outline level → HDSD template heading style display name."""
    return "A_HEADING 1" if level == 1 else f"A_Heading {level}"


def _match_fieldmap(heading_text: str, field_map: list[dict]) -> dict | None:
    """Prefix-match heading text (case-insensitive) against field_map entries."""
    needle = heading_text.strip().lower()
    # Prefer longer matches first (more specific)
    sorted_map = sorted(field_map, key=lambda e: -len(e.get("match", "")))
    for entry in sorted_map:
        key = entry.get("match", "").strip().lower()
        if not key:
            continue
        if needle.startswith(key) or key in needle:
            return entry
    return None


def _append_prose(doc, text: str, italic: bool = False):
    p = doc.add_paragraph(style="ETC_Content")
    run = p.add_run(text)
    if italic:
        run.italic = True
    return p


def _append_loop_table(doc, headers: list[str], loop: str, columns: list[str]):
    """Build a 4-row table: header / for-directive / data / endfor."""
    tbl = doc.add_table(rows=4, cols=len(headers))
    try:
        tbl.style = "Table Grid"
    except KeyError:
        pass
    hdr_cells = tbl.rows[0].cells
    for i, label in enumerate(headers):
        hdr_cells[i].text = label
        for r in hdr_cells[i].paragraphs[0].runs:
            r.bold = True
    tbl.rows[1].cells[0].text = f"{{%tr for {loop} %}}"
    data = tbl.rows[2].cells
    for i, expr in enumerate(columns):
        data[i].text = expr
    tbl.rows[3].cells[0].text = "{%tr endfor %}"
    return tbl


def build_body_from_outline(doc, outline: list[dict], field_map: list[dict]):
    """For each heading in outline: add heading paragraph, then insert content
    block according to field_map. Default fallback = `[CẦN BỔ SUNG: <heading>]`."""
    for node in outline:
        level = int(node["level"])
        text = node["text"].strip()
        if not text:
            continue
        # Heading
        doc.add_paragraph(text, style=_heading_style(level))
        # Content below (skip for container-only levels if desired — here always attempt)
        entry = _match_fieldmap(text, field_map)
        if entry is None:
            _append_prose(doc, f"[CẦN BỔ SUNG: {text}]", italic=True)
            continue
        kind = entry.get("kind", "placeholder")
        if kind == "prose":
            _append_prose(doc, entry["expr"])
        elif kind == "table":
            _append_loop_table(doc, entry["headers"], entry["loop"], entry["columns"])
        elif kind == "prose_with_diagram":
            _append_prose(doc, entry["expr"])
            _append_diagram(doc, entry["diagram"], entry.get("caption"))
        elif kind == "diagram_then_table":
            _append_diagram(doc, entry["diagram"], entry.get("caption"))
            _append_loop_table(doc, entry["headers"], entry["loop"], entry["columns"])
        elif kind == "diagram":
            _append_diagram(doc, entry["diagram"], entry.get("caption"))
        elif kind == "placeholder":
            _append_prose(doc, f"[CẦN BỔ SUNG: {text}]", italic=True)
        elif kind == "none":
            pass  # heading only (children will have content)
        else:
            _append_prose(doc, f"[CẦN BỔ SUNG: {text}]", italic=True)


def _append_diagram(doc, diagram_key: str, caption: str | None = None):
    """Insert a Jinja-guarded diagram block. Template uses `{{ diagram_<key> }}`
    resolved by render_docx.build_diagram_context() to InlineImage or None.
    If None → insert fallback text so doc still compiles."""
    # Guard with Jinja if (non-None) to skip missing diagrams cleanly
    doc.add_paragraph(f"{{%p if diagram_{diagram_key} %}}", style="ETC_Content")
    doc.add_paragraph(f"{{{{ diagram_{diagram_key} }}}}", style="ETC_Content")
    if caption:
        p = doc.add_paragraph(style="ETC_Content")
        r = p.add_run(caption)
        r.italic = True
    doc.add_paragraph("{%p else %}", style="ETC_Content")
    p = doc.add_paragraph(style="ETC_Content")
    r = p.add_run(f"[CẦN BỔ SUNG: Sơ đồ {diagram_key} — agent cần generate Mermaid source]")
    r.italic = True
    doc.add_paragraph("{%p endif %}", style="ETC_Content")


def _jinjafy_hdsd_base(source: Path, dest: Path, *, title_doc: str,
                       change_note: str, purpose_expr: str, scope_expr: str):
    """Shared cover/header/footer/preamble Jinjafication for TKKT + TKCS.
    Uses HDSD template as base — replaces HDSD title with doc-specific title."""
    dest.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(source, dest)
    doc = Document(dest)

    replacements = [
        ("<Tên dự án>", "{{ project.display_name }}"),
        ("<Mã dự án>", "{{ project.code }}"),
        ("Phiên bản: <xx>", "Phiên bản: {{ meta.version or '1.0' }}"),
        ("HỆ THỐNG QUẢN LÝ", "{{ project.display_name|upper }}"),
        ("TÀI LIỆU HƯỚNG DẪN SỬ DỤNG", title_doc.upper()),
        ("Tài liệu hướng dẫn sử dụng", title_doc),
        ("CÔNG TY CỔ PHẦN HỆ THỐNG CÔNG NGHỆ ETC", "{{ dev_unit|upper }}"),
        ("<Họ tên>", "[CẦN BỔ SUNG: Tên người ký]"),
        ("<Chức danh>", "[CẦN BỔ SUNG: Chức danh]"),
        ("<Màn hình 1>", ""),
        ("Ban hành: dd/mm/yyyy", "Ban hành: {{ meta.today }}"),
    ]
    replace_everywhere(doc, replacements)

    # Change history row
    if len(doc.tables) > 1:
        t = doc.tables[1]
        if len(t.rows) >= 2 and len(t.rows[1].cells) >= 5:
            c = t.rows[1].cells
            c[0].text = "{{ meta.today }}"
            c[1].text = "Toàn bộ tài liệu"
            c[2].text = "A"
            c[3].text = change_note
            c[4].text = "{{ meta.version or '1.0' }}"

    # Abbreviations + Related docs
    if len(doc.tables) > 4:
        jinjafy_loop_table(doc.tables[4], "overview.terms",
                           ["short", "full", "explanation"])
    if len(doc.tables) > 5:
        jinjafy_loop_table(doc.tables[5], "overview.references",
                           ["stt", "name", "ref"])

    # Section I
    fill_section_i_after_heading(doc, "Mục đích tài liệu", purpose_expr)
    fill_section_i_after_heading(doc, "Phạm vi tài liệu", scope_expr)

    # Clear Section II content (will be rebuilt from outline)
    clear_after_heading(doc, "NỘI DUNG")
    return doc


# ─────────────────────────── Template 2: TKKT ───────────────────────────

def jinjafy_tkkt(source: Path, dest: Path):
    """TKKT — outline theo QĐ 292/2025 (Khung KT CPĐT 4.0), 5 reference models."""
    doc = _jinjafy_hdsd_base(
        source, dest,
        title_doc="Tài liệu thiết kế kiến trúc",
        change_note="Tạo mới tài liệu thiết kế kiến trúc",
        purpose_expr="{{ architecture.purpose }}",
        scope_expr="{{ architecture.scope }}",
    )

    outline = _load_yaml(SCHEMA_DIR / "tkkt-outline.yaml")["outline"]
    fmap = _load_yaml(SCHEMA_DIR / "tkkt-fieldmap.yaml")["field_map"]
    build_body_from_outline(doc, outline, fmap)

    doc.save(dest)
    print(f"  TKKT saved: {dest} ({dest.stat().st_size // 1024} KB, {len(outline)} headings)")



# ─────────────────────────── Template 3: TKCS ───────────────────────────

def jinjafy_tkcs(source: Path, dest: Path):
    """TKCS — outline theo NĐ 45/2026 Điều 13 (11 H1 sections), fill all sections."""
    doc = _jinjafy_hdsd_base(
        source, dest,
        title_doc="Tài liệu thiết kế cơ sở",
        change_note="Tạo mới tài liệu thiết kế cơ sở",
        purpose_expr="{{ tkcs.purpose or architecture.purpose }}",
        scope_expr="{{ tkcs.scope or architecture.scope }}",
    )

    outline = _load_yaml(SCHEMA_DIR / "tkcs-outline.yaml")["outline"]
    fmap = _load_yaml(SCHEMA_DIR / "tkcs-fieldmap.yaml")["field_map"]
    build_body_from_outline(doc, outline, fmap)

    doc.save(dest)
    print(f"  TKCS saved: {dest} ({dest.stat().st_size // 1024} KB, {len(outline)} headings)")


# ─────────────────────────── CLI ───────────────────────────

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("kind", choices=["hdsd", "tkkt", "tkcs"])
    ap.add_argument("--source", required=True, help="Source template (original ETC or cloned HDSD)")
    ap.add_argument("--dest", default=None, help="Output path; defaults to templates/<name>.docx")
    args = ap.parse_args()

    source = Path(args.source)
    if not source.exists():
        print(f"ERROR: source not found: {source}", file=sys.stderr)
        sys.exit(1)

    default_names = {
        "hdsd": "huong-dan-su-dung.docx",
        "tkkt": "thiet-ke-kien-truc.docx",
        "tkcs": "thiet-ke-co-so.docx",
    }
    dest = Path(args.dest) if args.dest else (
        Path(__file__).parent.parent / "templates" / default_names[args.kind]
    )

    print(f"Jinjafying {args.kind.upper()}: {source} → {dest}")
    if args.kind == "hdsd":
        jinjafy_hdsd(source, dest)
    elif args.kind == "tkkt":
        jinjafy_tkkt(source, dest)
    elif args.kind == "tkcs":
        jinjafy_tkcs(source, dest)


if __name__ == "__main__":
    main()
