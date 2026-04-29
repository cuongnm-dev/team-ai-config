#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
extract_docx_schema.py — Analyze a Word template and emit a structural report.

Output JSON with:
  - Style catalog (all paragraph + character styles available)
  - Body element sequence (paragraphs with style + text, tables with dims)
  - Section/page settings (headers, footers, margins, orientation)
  - Detected placeholders (<...>, [Phần này...], [Ví dụ:...])
  - TOC element presence (SDT / field chars)
  - Cross-references & bookmarks (future use)

Purpose: human reviewer uses this as reference when forking a new template
(via tools/jinjafy_templates.py) to know which headings + tables need Jinja tags.
"""
from __future__ import annotations
import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Placeholder patterns commonly found in ETC templates
PLACEHOLDER_PATTERNS = [
    re.compile(r"<[^>]{1,60}>"),              # <Tên dự án>, <Họ tên>
    re.compile(r"^\s*\[Phần này.{0,200}\]"),   # [Phần này trình bày...]
    re.compile(r"\[Ví dụ:.{0,200}\]"),         # [Ví dụ: ...]
    re.compile(r"Ban hành:\s*dd/mm/yyyy"),     # Ban hành: dd/mm/yyyy
]


def iter_body_elements(doc):
    """Yield (index, kind, element) for each top-level body element.

    kind ∈ { 'p' (paragraph), 'tbl' (table), 'sectPr', 'sdt' (TOC content control) }
    """
    for i, child in enumerate(doc.element.body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        yield i, tag, child


def get_para_style(p_elem) -> str:
    ps = p_elem.find(f".//{{{WNS}}}pStyle")
    return ps.get(f"{{{WNS}}}val", "") if ps is not None else ""


def get_para_text(p_elem) -> str:
    return "".join(t.text or "" for t in p_elem.findall(f".//{{{WNS}}}t"))


def detect_placeholders(text: str) -> list[str]:
    hits = []
    for pat in PLACEHOLDER_PATTERNS:
        for m in pat.finditer(text):
            hits.append(m.group(0))
    return hits


def extract_styles(doc) -> dict:
    catalog = {"paragraph": [], "character": [], "table": []}
    for style in doc.styles:
        entry = {"name": style.name, "builtin": getattr(style, "builtin", False)}
        try:
            if style.type == 1:  # paragraph
                catalog["paragraph"].append(entry)
            elif style.type == 2:  # character
                catalog["character"].append(entry)
            elif style.type == 3:  # table
                catalog["table"].append(entry)
        except Exception:
            pass
    return catalog


def extract_tables_summary(doc) -> list[dict]:
    out = []
    for i, tbl in enumerate(doc.tables):
        rows = len(tbl.rows)
        cols = len(tbl.rows[0].cells) if rows else 0
        first_row_texts = []
        if rows:
            for c in tbl.rows[0].cells:
                first_row_texts.append(c.text.strip()[:60])
        out.append({
            "index": i,
            "rows": rows,
            "cols": cols,
            "first_row": first_row_texts,
        })
    return out


def extract_body_sequence(doc) -> list[dict]:
    """Walk body top-level elements with style + text + placeholder flag."""
    seq = []
    tbl_counter = 0
    for i, kind, elem in iter_body_elements(doc):
        if kind == "p":
            style = get_para_style(elem)
            text = get_para_text(elem)
            ph = detect_placeholders(text)
            seq.append({
                "idx": i,
                "kind": "p",
                "style": style,
                "text": text[:300],
                "placeholders": ph,
                "is_heading": "HEADING" in style.upper() or "Heading" in style,
            })
        elif kind == "tbl":
            seq.append({
                "idx": i,
                "kind": "tbl",
                "table_index": tbl_counter,
            })
            tbl_counter += 1
        elif kind == "sdt":
            # Structured Document Tag — usually TOC
            seq.append({"idx": i, "kind": "sdt"})
        elif kind == "sectPr":
            seq.append({"idx": i, "kind": "sectPr"})
        else:
            seq.append({"idx": i, "kind": kind})
    return seq


def extract_headers_footers(doc) -> dict:
    out = {"sections": []}
    for s_idx, section in enumerate(doc.sections):
        sec = {
            "section_index": s_idx,
            "different_first_page": section.different_first_page_header_footer,
            "headers": {},
            "footers": {},
            "margins_inch": {
                "top": float(section.top_margin.inches) if section.top_margin else None,
                "bottom": float(section.bottom_margin.inches) if section.bottom_margin else None,
                "left": float(section.left_margin.inches) if section.left_margin else None,
                "right": float(section.right_margin.inches) if section.right_margin else None,
            },
        }
        for name, part in [
            ("default", section.header),
            ("first_page", section.first_page_header),
            ("even_page", section.even_page_header),
        ]:
            if part is None:
                continue
            texts = [p.text for p in part.paragraphs if p.text.strip()]
            if texts:
                sec["headers"][name] = texts
        for name, part in [
            ("default", section.footer),
            ("first_page", section.first_page_footer),
            ("even_page", section.even_page_footer),
        ]:
            if part is None:
                continue
            texts = [p.text for p in part.paragraphs if p.text.strip()]
            if texts:
                sec["footers"][name] = texts
        out["sections"].append(sec)
    return out


def extract_workbook(path: Path) -> dict:
    doc = Document(path)
    body_seq = extract_body_sequence(doc)
    heading_only = [e for e in body_seq if e.get("is_heading")]

    placeholder_map: dict[str, list[int]] = {}
    for e in body_seq:
        for ph in e.get("placeholders", []):
            placeholder_map.setdefault(ph, []).append(e["idx"])

    return {
        "template_path": str(path),
        "total_body_elements": len(body_seq),
        "paragraph_count": len([e for e in body_seq if e["kind"] == "p"]),
        "table_count": len([e for e in body_seq if e["kind"] == "tbl"]),
        "heading_count": len(heading_only),
        "styles": extract_styles(doc),
        "headings": heading_only,
        "tables_summary": extract_tables_summary(doc),
        "body_sequence": body_seq,
        "placeholders_found": placeholder_map,
        "headers_footers": extract_headers_footers(doc),
    }


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("template")
    ap.add_argument("--out", default=None)
    args = ap.parse_args()
    path = Path(args.template)
    if not path.exists():
        print(f"ERROR: not found: {path}", file=sys.stderr)
        sys.exit(1)
    rep = extract_workbook(path)
    txt = json.dumps(rep, ensure_ascii=False, indent=2)
    if args.out:
        Path(args.out).write_text(txt, encoding="utf-8")
        print(f"Wrote {args.out}  ({rep['paragraph_count']} paras, "
              f"{rep['table_count']} tables, {rep['heading_count']} headings, "
              f"{len(rep['placeholders_found'])} distinct placeholders)")
    else:
        print(txt)


if __name__ == "__main__":
    main()
